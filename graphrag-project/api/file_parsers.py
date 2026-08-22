import csv as csv_module
import json
import logging
import os
import re
import subprocess
import tempfile
from pathlib import Path
from typing import Any, Dict, List, Optional

import pdfplumber
from docx import Document as DocxDocument

from config import settings
import structural_extraction

logger = logging.getLogger("graphrag")

# OCR is optional: only pytesseract + pdf2image import failures disable it,
# they don't break the module for everyone else. pdf2image also needs the
# poppler-utils system binaries (pdftoppm/pdftocairo) on PATH - if those are
# missing, OCR calls fail at runtime instead of import time, so _extract_pdf
# catches that separately too.
try:
    import pytesseract
    from pdf2image import convert_from_path

    _OCR_AVAILABLE = True
except ImportError:
    _OCR_AVAILABLE = False

SUPPORTED_EXTENSIONS = {
    ".pdf", ".json", ".txt", ".log", ".md", ".doc", ".docx",
    ".csv", ".tsv", ".xlsx", ".html", ".htm", ".xml", ".eml",
}

_CSV_ROWS_PER_BLOCK = 20


def extract_text(path: str, filename: str) -> List[Dict[str, Any]]:
    """
    Returns a list of {"text": str, "metadata": dict} blocks rather than a
    single string. A block is a logical unit - a PDF page, a docx
    section, a batch of spreadsheet rows, an XML/JSON record - and
    ingest.py's chunker splits each block independently, so a chunk never
    silently spans two unrelated pages/sections. Useful metadata (page
    number, section heading, sheet name, row range) rides along into the
    chunk payload for citation purposes.
    """
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return _extract_pdf(path)
    if ext == ".json":
        return _extract_json(path)
    if ext in (".txt", ".log", ".md"):
        return _extract_plain_text(path)
    if ext == ".docx":
        return _extract_docx(path)
    if ext == ".doc":
        return _extract_doc_legacy(path)
    if ext == ".csv":
        return _extract_csv(path, delimiter=",")
    if ext == ".tsv":
        return _extract_csv(path, delimiter="\t")
    if ext == ".xlsx":
        return _extract_xlsx(path)
    if ext in (".html", ".htm"):
        return _extract_html(path)
    if ext == ".xml":
        return _extract_xml(path)
    if ext == ".eml":
        return _extract_eml(path)
    raise ValueError(f"Unsupported file type: {ext}")


def _extract_plain_text(path: str) -> List[Dict[str, Any]]:
    with open(path, "rb") as f:
        raw = f.read()
    text = None
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            text = raw.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    if text is None:
        text = raw.decode("utf-8", errors="ignore")
    return [{"text": text, "metadata": {}}]


# ---------------------------------------------------------------------------
# PDF (with OCR fallback for scanned pages)
# ---------------------------------------------------------------------------

def _extract_pdf(path: str) -> List[Dict[str, Any]]:
    blocks: List[Dict[str, Any]] = []
    ocr_needed_pages: List[int] = []

    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages):
            text = (page.extract_text() or "").strip()
            if text:
                blocks.append({"text": text, "metadata": {"page": i + 1}})
            else:
                ocr_needed_pages.append(i + 1)
            # Releases pdfplumber's cached parsed objects for this page -
            # without this, memory for a large PDF grows for the whole
            # document even though we only need one page's text at a time.
            page.flush_cache()

    if ocr_needed_pages:
        if settings.OCR_ENABLED and _OCR_AVAILABLE:
            ocr_blocks = _ocr_pdf_pages(path, ocr_needed_pages)
            blocks.extend(ocr_blocks)
            blocks.sort(key=lambda b: b["metadata"].get("page", 0))
        else:
            reason = "disabled (OCR_ENABLED=false)" if not settings.OCR_ENABLED else "unavailable (pytesseract/pdf2image not installed)"
            logger.warning(
                f"file_parsers: {len(ocr_needed_pages)} page(s) had no extractable "
                f"text and OCR is {reason} - those pages will be skipped."
            )

    return blocks


def _ocr_pdf_pages(path: str, page_numbers: List[int]) -> List[Dict[str, Any]]:
    """OCRs specific 1-indexed pages of a PDF (only called for pages whose
    native text extraction returned nothing). Renders each page to an
    image at OCR_DPI and runs Tesseract on it. A failure on one page is
    logged and skipped rather than failing the whole document - a
    partially-OCR'd document is still more useful than none at all."""
    blocks = []
    for page_num in page_numbers:
        try:
            images = convert_from_path(path, first_page=page_num, last_page=page_num, dpi=settings.OCR_DPI)
            if not images:
                continue
            text = pytesseract.image_to_string(images[0]).strip()
            if text:
                blocks.append({"text": text, "metadata": {"page": page_num, "ocr": True}})
        except Exception:
            logger.exception(f"file_parsers: OCR failed for page {page_num} of {path}")
    return blocks


# ---------------------------------------------------------------------------
# DOCX / legacy DOC
# ---------------------------------------------------------------------------

def _extract_docx(path: str) -> List[Dict[str, Any]]:
    doc = DocxDocument(path)
    blocks: List[Dict[str, Any]] = []
    current_heading = None
    current_lines: List[str] = []

    def _flush():
        text = "\n".join(current_lines).strip()
        if text:
            metadata = {"section": current_heading} if current_heading else {}
            blocks.append({"text": text, "metadata": metadata})

    for p in doc.paragraphs:
        style_name = (p.style.name if p.style else "") or ""
        text = p.text.strip()
        if not text:
            continue
        if style_name.lower().startswith("heading") or style_name.lower() == "title":
            _flush()
            current_lines = []
            current_heading = text
        else:
            current_lines.append(text)
    _flush()

    for i, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            blocks.append({"text": "\n".join(rows), "metadata": {"table": i + 1}})

    return blocks


def _extract_doc_legacy(path: str) -> List[Dict[str, Any]]:
    """Legacy binary .doc via headless LibreOffice conversion (no pure-Python
    parser handles the old binary format reliably). Timeout scales with
    file size instead of a single fixed value, since a large legacy .doc
    can plausibly take longer to convert than a small one."""
    file_size_mb = os.path.getsize(path) / (1024 * 1024)
    timeout = min(
        settings.DOC_CONVERT_TIMEOUT_BASE + int(file_size_mb * settings.DOC_CONVERT_TIMEOUT_PER_MB),
        settings.DOC_CONVERT_TIMEOUT_MAX,
    )
    with tempfile.TemporaryDirectory() as out_dir, tempfile.TemporaryDirectory() as profile_dir:
        result = subprocess.run(
            [
                "soffice",
                "--headless",
                "--norestore",
                f"-env:UserInstallation=file://{profile_dir}",
                "--convert-to", "txt:Text",
                "--outdir", out_dir,
                path,
            ],
            capture_output=True,
            timeout=timeout,
        )
        expected = Path(out_dir) / (Path(path).stem + ".txt")
        if not expected.exists():
            stderr = result.stderr.decode(errors="ignore")[:500]
            raise RuntimeError(f"Failed to convert .doc file via LibreOffice: {stderr}")
        text = expected.read_text(encoding="utf-8", errors="ignore")
    return [{"text": text, "metadata": {}}]


# ---------------------------------------------------------------------------
# JSON (streaming for large files) / XML
# ---------------------------------------------------------------------------

def _flatten_json(obj, path: str = "") -> List[str]:
    lines: List[str] = []
    if isinstance(obj, dict):
        for k, v in obj.items():
            lines.extend(_flatten_json(v, f"{path}.{k}" if path else str(k)))
    elif isinstance(obj, list):
        for i, item in enumerate(obj):
            lines.extend(_flatten_json(item, f"{path}[{i}]"))
    else:
        lines.append(f"{path}: {obj}" if path else str(obj))
    return lines


def _json_block_metadata(record_index: Optional[int], value: Any) -> Dict[str, Any]:
    """Builds a block's metadata dict, attaching a precomputed structural
    entity/relationship extraction when enabled - ingest.py checks for
    this and, when present, skips the LLM extraction call entirely for
    chunks from this block (see structural_extraction.py for why and
    ingest.py's _extract_batch_concurrently for how)."""
    metadata: Dict[str, Any] = {}
    if record_index is not None:
        metadata["record"] = record_index
    if settings.JSON_STRUCTURAL_EXTRACTION_ENABLED:
        metadata["structural_extraction"] = structural_extraction.extract_from_json_value(value)
    return metadata


def _extract_json(path: str) -> List[Dict[str, Any]]:
    size_mb = os.path.getsize(path) / (1024 * 1024)
    if size_mb > settings.JSON_STREAMING_THRESHOLD_MB:
        try:
            return _extract_json_streaming(path)
        except Exception as e:
            logger.warning(
                f"file_parsers: streaming JSON parse failed ({e}), "
                f"falling back to full in-memory parse for {path}"
            )

    with open(path, "r", encoding="utf-8") as f:
        data = json.load(f)

    # A top-level list of records is treated as one block per record so
    # the chunker never mixes two unrelated records into the same chunk.
    if isinstance(data, list):
        blocks = []
        for i, item in enumerate(data):
            flat = _flatten_json(item)
            if flat:
                blocks.append({"text": "\n".join(flat), "metadata": _json_block_metadata(i, item)})
        return blocks

    flat = _flatten_json(data)
    return [{"text": "\n".join(flat), "metadata": _json_block_metadata(None, data)}] if flat else []


def _extract_json_streaming(path: str) -> List[Dict[str, Any]]:
    """
    For large JSON files, streams a top-level array with ijson instead of
    loading the entire parsed structure into memory at once - only used
    above JSON_STREAMING_THRESHOLD_MB, since ijson has real per-item
    overhead not worth paying for small files. Only handles the (very
    common) case of a top-level JSON array; a large top-level object
    raises, which the caller catches and falls back to the in-memory path
    for.
    """
    import ijson

    blocks = []
    with open(path, "rb") as f:
        for i, item in enumerate(ijson.items(f, "item")):
            flat = _flatten_json(item)
            if flat:
                blocks.append({"text": "\n".join(flat), "metadata": _json_block_metadata(i, item)})
    if not blocks:
        raise ValueError("no top-level array items found (streaming path only handles JSON arrays)")
    return blocks


def _extract_xml(path: str) -> List[Dict[str, Any]]:
    import xml.etree.ElementTree as ET

    tree = ET.parse(path)
    root = tree.getroot()

    def _flatten_xml(el, path_str: str = "") -> List[str]:
        lines: List[str] = []
        tag = el.tag.split("}")[-1]  # strip XML namespace prefix if present
        current_path = f"{path_str}.{tag}" if path_str else tag
        for k, v in el.attrib.items():
            lines.append(f"{current_path}[@{k}]: {v}")
        text = (el.text or "").strip()
        children = list(el)
        if text and not children:
            lines.append(f"{current_path}: {text}")
        for child in children:
            lines.extend(_flatten_xml(child, current_path))
        return lines

    # Each direct child of the root becomes one block, mirroring how a
    # top-level JSON array is chunked one record per block - keeps a
    # chunk from spanning unrelated XML records.
    children = list(root)
    if children:
        blocks = []
        for i, child in enumerate(children):
            flat = _flatten_xml(child)
            if flat:
                blocks.append({"text": "\n".join(flat), "metadata": {"record": i}})
        return blocks

    flat = _flatten_xml(root)
    return [{"text": "\n".join(flat), "metadata": {}}] if flat else []


# ---------------------------------------------------------------------------
# CSV / TSV / XLSX
# ---------------------------------------------------------------------------

def _extract_csv(path: str, delimiter: str = ",") -> List[Dict[str, Any]]:
    """Batches rows into blocks of _CSV_ROWS_PER_BLOCK (formatted as
    "header: value" pairs, like the JSON flattener) rather than chunking
    the raw CSV text by character count, which would slice rows in half
    at arbitrary points."""
    with open(path, "r", newline="", encoding="utf-8", errors="ignore") as f:
        reader = csv_module.reader(f, delimiter=delimiter)
        try:
            header = next(reader)
        except StopIteration:
            return []

        blocks = []
        batch_lines: List[str] = []
        batch_start = 1
        row_num = 1
        for row in reader:
            row_num += 1
            pairs = [f"{h}: {v}" for h, v in zip(header, row) if h]
            if not pairs:
                continue
            batch_lines.append("; ".join(pairs))
            if len(batch_lines) >= _CSV_ROWS_PER_BLOCK:
                blocks.append({"text": "\n".join(batch_lines), "metadata": {"rows": f"{batch_start}-{row_num}"}})
                batch_lines = []
                batch_start = row_num + 1
        if batch_lines:
            blocks.append({"text": "\n".join(batch_lines), "metadata": {"rows": f"{batch_start}-{row_num}"}})
        return blocks


def _extract_xlsx(path: str) -> List[Dict[str, Any]]:
    import openpyxl

    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    blocks: List[Dict[str, Any]] = []
    try:
        for sheet in wb.worksheets:
            rows_iter = sheet.iter_rows(values_only=True)
            try:
                header = next(rows_iter)
            except StopIteration:
                continue
            header = [str(h) if h is not None else "" for h in header]

            batch_lines: List[str] = []
            batch_start = 1
            row_num = 1
            for row in rows_iter:
                row_num += 1
                if row is None or all(v is None for v in row):
                    continue
                pairs = [f"{h}: {v}" for h, v in zip(header, row) if h]
                if not pairs:
                    continue
                batch_lines.append("; ".join(pairs))
                if len(batch_lines) >= _CSV_ROWS_PER_BLOCK:
                    blocks.append({
                        "text": "\n".join(batch_lines),
                        "metadata": {"sheet": sheet.title, "rows": f"{batch_start}-{row_num}"},
                    })
                    batch_lines = []
                    batch_start = row_num + 1
            if batch_lines:
                blocks.append({
                    "text": "\n".join(batch_lines),
                    "metadata": {"sheet": sheet.title, "rows": f"{batch_start}-{row_num}"},
                })
    finally:
        wb.close()
    return blocks


# ---------------------------------------------------------------------------
# HTML
# ---------------------------------------------------------------------------

def _extract_html(path: str) -> List[Dict[str, Any]]:
    from bs4 import BeautifulSoup

    with open(path, "rb") as f:
        raw = f.read()
    soup = BeautifulSoup(raw, "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()

    blocks: List[Dict[str, Any]] = []

    # Tables are extracted (and removed from the tree) first, so their
    # cell text doesn't also get swept up as loose paragraph text in the
    # pass below, and so each table stays one coherent block rather than
    # scattering cell text across unrelated section blocks.
    for i, table in enumerate(soup.find_all("table")):
        rows = []
        for tr in table.find_all("tr"):
            cells = [c.get_text(" ", strip=True) for c in tr.find_all(["td", "th"])]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            blocks.append({"text": "\n".join(rows), "metadata": {"table": i + 1}})
        table.decompose()

    current_heading = None
    current_lines: List[str] = []

    def _flush():
        text = "\n".join(current_lines).strip()
        if text:
            metadata = {"section": current_heading} if current_heading else {}
            blocks.append({"text": text, "metadata": metadata})

    body = soup.body or soup
    for el in body.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li"]):
        if re.match(r"^h[1-6]$", el.name):
            heading_text = el.get_text(strip=True)
            if heading_text:
                _flush()
                current_lines = []
                current_heading = heading_text
        else:
            text = el.get_text(" ", strip=True)
            if text:
                current_lines.append(text)
    _flush()

    if not blocks:
        # No structural tags matched (e.g. a page built entirely from divs)
        # - fall back to grabbing all visible text as one block rather than
        # returning nothing.
        text = soup.get_text("\n", strip=True)
        if text:
            blocks = [{"text": text, "metadata": {}}]

    return blocks


# ---------------------------------------------------------------------------
# EML (email)
# ---------------------------------------------------------------------------

def _extract_eml(path: str) -> List[Dict[str, Any]]:
    import email
    from email import policy as email_policy

    with open(path, "rb") as f:
        msg = email.message_from_binary_file(f, policy=email_policy.default)

    header_lines = []
    for header in ("From", "To", "Cc", "Subject", "Date"):
        value = msg.get(header)
        if value:
            header_lines.append(f"{header}: {value}")

    body_text = ""
    if msg.is_multipart():
        for part in msg.walk():
            if part.get_content_type() == "text/plain":
                try:
                    body_text = part.get_content()
                    break
                except Exception:
                    continue
        if not body_text:
            for part in msg.walk():
                if part.get_content_type() == "text/html":
                    try:
                        body_text = _html_to_text(part.get_content())
                        break
                    except Exception:
                        continue
    else:
        try:
            body_text = msg.get_content()
        except Exception:
            body_text = ""
        if msg.get_content_type() == "text/html" and body_text:
            body_text = _html_to_text(body_text)

    blocks = []
    if header_lines:
        blocks.append({"text": "\n".join(header_lines), "metadata": {"section": "headers"}})
    if body_text and body_text.strip():
        blocks.append({"text": body_text.strip(), "metadata": {"section": "body"}})
    return blocks


def _html_to_text(html: str) -> str:
    from bs4 import BeautifulSoup

    return BeautifulSoup(html, "html.parser").get_text("\n", strip=True)
